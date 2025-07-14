import json
import os
import sys
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import jsonschema
from .style import STYLE_CONFIG

class CVData:
    """Handles loading, validation, and access to CV data."""
    def __init__(self, json_path, schema_path):
        self.json_path = json_path
        self.schema_path = schema_path
        self.data = self._load_and_validate_data()

    def _load_and_validate_data(self):
        """Loads the JSON data and validates it against the schema."""
        try:
            with open(self.json_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            with open(self.schema_path, 'r', encoding='utf-8') as file:
                schema = json.load(file)
            
            jsonschema.validate(instance=data, schema=schema)
            return data
        except FileNotFoundError as e:
            print(f"Error: {e.filename} not found.")
            sys.exit(1)
        except jsonschema.exceptions.ValidationError as e:
            print(f"JSON validation error: {e.message}")
            sys.exit(1)
        except json.JSONDecodeError:
            print(f"Error decoding JSON from {self.json_path}")
            sys.exit(1)

class DocxGenerator:
    """Generates the CV document from CVData."""
    def __init__(self, cv_data, style_config):
        self.cv_data = cv_data
        self.style = style_config
        self.doc = Document()
        self._apply_document_styles()

    def _apply_document_styles(self):
        """Applies base styles and margins to the document."""
        for section in self.doc.sections:
            section.top_margin = self.style["margins"]["top"]
            section.bottom_margin = self.style["margins"]["bottom"]
            section.left_margin = self.style["margins"]["left"]
            section.right_margin = self.style["margins"]["right"]

        style = self.doc.styles['Normal']
        style.font.name = self.style["font_name"]
        style.font.size = Pt(self.style["font_size"])
        
        self._create_custom_bullet_style()

    def _create_custom_bullet_style(self):
        """Creates the custom bullet style."""
        bullet_style_def = self.style["bullet_style"]
        list_style = self.doc.styles.add_style('CustomBulletStyle', WD_STYLE_TYPE.PARAGRAPH)
        list_style.base_style = self.doc.styles['List Bullet']
        list_style.paragraph_format.left_indent = bullet_style_def["left_indent"]
        list_style.paragraph_format.first_line_indent = bullet_style_def["first_line_indent"]
        list_style.paragraph_format.line_spacing = bullet_style_def["line_spacing"]
        list_style.paragraph_format.space_after = bullet_style_def["space_after"]
        list_style.paragraph_format.space_before = bullet_style_def["space_before"]
        list_style.font.name = self.style["font_name"]
        list_style.font.size = Pt(self.style["font_size"])

    def generate(self, output_path):
        """Generates and saves the full CV document."""
        self._create_personal_info_section()
        self._create_summary_section()
        self._create_experience_section()
        self._create_education_section()
        self._create_projects_section()
        self._create_skills_section()
        self._create_certifications_section()
        self._create_publications_section()
        self._create_awards_section()
        self._create_volunteer_section()
        self._create_languages_section()
        self._create_references_section()
        self.doc.save(output_path)
        print(f"CV saved as {output_path}")

    def _create_personal_info_section(self):
        personal_info = self.cv_data.data.get("personalInfo", {})
        if not personal_info:
            return

        # Full Name
        name_parts = [personal_info.get(key) for key in ["firstName", "middleName", "lastName"] if personal_info.get(key)]
        full_name = " ".join(name_parts)
        if full_name:
            name_para = self.doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(full_name)
            name_run.font.size = Pt(self.style["name_font_size"])
            name_run.font.name = self.style["font_name"]
            name_run.font.bold = True
            self._set_paragraph_spacing(name_para, after=3)

        # Contact Info Rows
        self._add_contact_row([
            {"text": personal_info.get("email"), "url": f"mailto:{personal_info.get('email')}"},
            {"text": personal_info.get("phone"), "url": f"tel:{personal_info.get('phone')}"}
        ])
        
        contact_row2_items = []
        if personal_info.get("linkedIn"):
            contact_row2_items.append({"text": self._format_url(personal_info["linkedIn"]), "url": personal_info["linkedIn"]})
        if personal_info.get("githubUrl"):
            contact_row2_items.append({"text": self._format_url(personal_info["githubUrl"]), "url": personal_info["githubUrl"]})
        for field in ["website", "portfolio", "blog"]:
            if personal_info.get(field):
                contact_row2_items.append({"text": self._extract_domain(personal_info[field]), "url": personal_info[field]})
        self._add_contact_row(contact_row2_items, after=3)

    def _create_summary_section(self):
        """Creates the professional summary section."""
        summary = self.cv_data.data.get("professionalSummary")
        if not summary:
            return
        
        self._add_section_header("PROFESSIONAL SUMMARY")
        summary_para = self.doc.add_paragraph()
        self._set_paragraph_spacing(summary_para, after=3)
        summary_run = summary_para.add_run(summary)
        self._set_font_properties(summary_run)

    def _add_contact_row(self, items, after=2):
        items = [item for item in items if item.get("text")]
        if not items:
            return
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_spacing(para, after=after)
        
        for i, item in enumerate(items):
            if i > 0:
                para.add_run(" | ")
            self._add_hyperlink(para, item["url"], item["text"])
        
        for run in para.runs:
            self._set_font_properties(run)

    def _create_experience_section(self):
        work_experience = self.cv_data.data.get("workExperience", [])
        if not work_experience:
            return
        
        self._add_section_header("EXPERIENCE")
        for job in work_experience:
            if job.get("company"):
                company_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(company_para)
                company_run = company_para.add_run(job["company"])
                self._set_font_properties(company_run, bold=True)
                
                # Add location if available
                if job.get("location"):
                    company_para.add_run(f" | {job['location']}")
                    for run in company_para.runs[-1:]:
                        self._set_font_properties(run, bold=True)
                
                company_url = self._find_company_url(job)
                if company_url:
                    company_para.add_run(" (")
                    self._add_hyperlink(company_para, company_url, self._extract_domain(company_url), underline=False)
                    company_para.add_run(")")
                    for run in company_para.runs:
                        self._set_font_properties(run, bold=True)

            if job.get("position"):
                position_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(position_para)
                position_run = position_para.add_run(job["position"])
                self._set_font_properties(position_run, bold=True)

            start_date = self._format_date(job.get("startDate"))
            end_date = "Present" if job.get("current") else self._format_date(job.get("endDate"))
            if start_date or end_date:
                date_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(date_para)
                date_range = f"{start_date} - {end_date}" if start_date and end_date else start_date or end_date
                date_run = date_para.add_run(date_range)
                self._set_font_properties(date_run)

            if job.get("description"):
                self._add_bullet_point(job["description"])

            for achievement in job.get("achievements", []):
                if achievement:
                    self._add_bullet_point(achievement if achievement.endswith('.') else f"{achievement}.")
            
            # Add technologies if available
            if job.get("technologies"):
                tech_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(tech_para, after=3)
                tech_run = tech_para.add_run("Technologies used: ")
                self._set_font_properties(tech_run, bold=True)
                tech_list_run = tech_para.add_run(", ".join(job["technologies"]))
                self._set_font_properties(tech_list_run)

    def _create_education_section(self):
        education = self.cv_data.data.get("education", [])
        if not education:
            return

        self._add_section_header("EDUCATION")
        for edu in education:
            if edu.get("institution"):
                inst_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(inst_para)
                inst_run = inst_para.add_run(edu["institution"])
                self._set_font_properties(inst_run)

            degree_parts = [edu.get("degree"), edu.get("field")]
            if any(degree_parts):
                degree_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(degree_para)
                degree_run = degree_para.add_run(" | ".join(filter(None, degree_parts)))
                self._set_font_properties(degree_run)
                
                # Add GPA if available
                if edu.get("gpa"):
                    degree_para.add_run(f" | GPA: {edu['gpa']:.2f}/4.0")
                    for run in degree_para.runs[-1:]:
                        self._set_font_properties(run)

            loc_grad_parts = [edu.get("location")]
            if edu.get("graduationDate"):
                loc_grad_parts.append(f"Graduated in {self._format_date(edu['graduationDate'])}")
            if any(loc_grad_parts):
                loc_date_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(loc_date_para, after=3)
                loc_date_run = loc_date_para.add_run(" | ".join(filter(None, loc_grad_parts)))
                self._set_font_properties(loc_date_run)
            
            # Add honors if available
            for honor in edu.get("honors", []):
                if honor:
                    self._add_bullet_point(honor)
            
            # Add relevant courses if available
            if edu.get("relevantCourses"):
                courses_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(courses_para, after=3)
                courses_run = courses_para.add_run("Relevant Courses: ")
                self._set_font_properties(courses_run, bold=True)
                courses_list_run = courses_para.add_run(", ".join(edu["relevantCourses"]))
                self._set_font_properties(courses_list_run)

    def _create_projects_section(self):
        """Creates the projects section."""
        projects = self.cv_data.data.get("projects", [])
        if not projects:
            return
        
        self._add_section_header("PROJECTS")
        for project in projects:
            # Project name with optional URL
            if project.get("name"):
                project_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(project_para)
                
                if project.get("url"):
                    self._add_hyperlink(project_para, project["url"], project["name"])
                    # Make the hyperlink bold
                    for run in project_para.runs:
                        self._set_font_properties(run, bold=True)
                else:
                    project_run = project_para.add_run(project["name"])
                    self._set_font_properties(project_run, bold=True)
            
            # Dates
            start_date = self._format_date(project.get("startDate"))
            end_date = self._format_date(project.get("endDate"))
            if start_date or end_date:
                date_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(date_para)
                date_range = f"{start_date} - {end_date}" if start_date and end_date else start_date or end_date or ""
                date_run = date_para.add_run(date_range)
                self._set_font_properties(date_run)
            
            # Description
            if project.get("description"):
                desc_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(desc_para)
                desc_run = desc_para.add_run(project["description"])
                self._set_font_properties(desc_run)
            
            # Highlights
            for highlight in project.get("highlights", []):
                if highlight:
                    self._add_bullet_point(highlight if highlight.endswith('.') else f"{highlight}.")
            
            # Technologies
            if project.get("technologies"):
                tech_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(tech_para, after=3)
                tech_run = tech_para.add_run("Technologies: ")
                self._set_font_properties(tech_run, bold=True)
                tech_list_run = tech_para.add_run(", ".join(project["technologies"]))
                self._set_font_properties(tech_list_run)

    def _create_languages_section(self):
        languages = self.cv_data.data.get("languages", [])
        if not languages:
            return
        
        self._add_section_header("LANGUAGES")
        lang_entries = [self._format_language_entry(lang) for lang in languages]
        if lang_entries:
            lang_para = self.doc.add_paragraph()
            self._set_paragraph_spacing(lang_para, after=3)
            lang_run = lang_para.add_run(", ".join(filter(None, lang_entries)))
            self._set_font_properties(lang_run)

    def _create_skills_section(self):
        skills = self.cv_data.data.get("skills", [])
        if not skills:
            return
        
        # Check if skills have categories
        has_categories = any(isinstance(skill, dict) and skill.get("category") for skill in skills)
        
        self._add_section_header("SKILLS" if has_categories else "TECHNOLOGIES")
        
        if has_categories:
            # Group skills by category
            from collections import defaultdict
            categories = defaultdict(list)
            uncategorized = []
            
            for skill in skills:
                if isinstance(skill, dict):
                    category = skill.get("category", "Other")
                    skill_info = skill.get("name", "")
                    if skill.get("level"):
                        skill_info += f" ({skill['level']})"
                    if category:
                        categories[category].append(skill_info)
                    else:
                        uncategorized.append(skill_info)
                elif isinstance(skill, str):
                    uncategorized.append(skill)
            
            # Render categorized skills
            for category, skill_list in categories.items():
                if skill_list:
                    cat_para = self.doc.add_paragraph()
                    self._set_paragraph_spacing(cat_para)
                    cat_run = cat_para.add_run(f"{category}: ")
                    self._set_font_properties(cat_run, bold=True)
                    skills_run = cat_para.add_run(", ".join(skill_list))
                    self._set_font_properties(skills_run)
            
            # Render uncategorized skills
            if uncategorized:
                uncat_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(uncat_para)
                uncat_run = uncat_para.add_run("Other: " if categories else "")
                if categories:
                    self._set_font_properties(uncat_run, bold=True)
                skills_run = uncat_para.add_run(", ".join(uncategorized))
                self._set_font_properties(skills_run)
        else:
            # Simple list format (existing behavior)
            skill_names = []
            for skill in skills:
                if isinstance(skill, dict):
                    name = skill.get("name")
                    if name:
                        if skill.get("level"):
                            skill_names.append(f"{name} ({skill['level']})")
                        else:
                            skill_names.append(name)
                elif isinstance(skill, str):
                    skill_names.append(skill)
            
            skill_names = list(filter(None, skill_names))
            if skill_names:
                skills_text = ", ".join(skill_names)
                if not skills_text.endswith("."):
                    skills_text += "."
                
                tech_para = self.doc.add_paragraph()
                tech_run = tech_para.add_run(skills_text)
                self._set_font_properties(tech_run)

    def _create_certifications_section(self):
        """Creates the certifications section."""
        certifications = self.cv_data.data.get("certifications", [])
        if not certifications:
            return
        
        self._add_section_header("CERTIFICATIONS")
        for cert in certifications:
            cert_para = self.doc.add_paragraph()
            self._set_paragraph_spacing(cert_para, after=3)
            
            # Certification name (bold)
            if cert.get("name"):
                cert_run = cert_para.add_run(cert["name"])
                self._set_font_properties(cert_run, bold=True)
            
            # Issuer
            if cert.get("issuer"):
                cert_para.add_run(f", {cert['issuer']}")
                for run in cert_para.runs[-1:]:
                    self._set_font_properties(run)
            
            # Date obtained
            if cert.get("dateObtained"):
                date_str = self._format_date(cert["dateObtained"])
                cert_para.add_run(f", Issued {date_str}")
                for run in cert_para.runs[-1:]:
                    self._set_font_properties(run)
            
            # Expiry date
            if cert.get("expiryDate"):
                expiry_str = self._format_date(cert["expiryDate"])
                cert_para.add_run(f" (Expires {expiry_str})")
                for run in cert_para.runs[-1:]:
                    self._set_font_properties(run)
            
            # Credential URL
            if cert.get("credentialUrl"):
                cert_para.add_run(" ")
                self._add_hyperlink(cert_para, cert["credentialUrl"], "[View Certificate]")
                for run in cert_para.runs:
                    if run.text and run.text not in ["[View Certificate]"]:
                        self._set_font_properties(run)

    def _create_publications_section(self):
        """Creates the publications section."""
        publications = self.cv_data.data.get("publications", [])
        if not publications:
            return
        
        self._add_section_header("PUBLICATIONS")
        for pub in publications:
            pub_para = self.doc.add_paragraph()
            self._set_paragraph_spacing(pub_para, after=3)
            
            # Title (bold)
            if pub.get("title"):
                title_run = pub_para.add_run(pub["title"])
                self._set_font_properties(title_run, bold=True)
            
            # Authors
            if pub.get("authors") and isinstance(pub["authors"], list):
                authors_str = ", ".join(pub["authors"])
                pub_para.add_run(f". {authors_str}")
                for run in pub_para.runs[-1:]:
                    self._set_font_properties(run)
            
            # Publisher
            if pub.get("publisher"):
                pub_para.add_run(f". {pub['publisher']}")
                for run in pub_para.runs[-1:]:
                    self._set_font_properties(run)
            
            # Date
            if pub.get("date"):
                date_str = self._format_date(pub["date"])
                pub_para.add_run(f", {date_str}")
                for run in pub_para.runs[-1:]:
                    self._set_font_properties(run)
            
            # DOI or URL
            if pub.get("doi"):
                pub_para.add_run(" ")
                doi_url = f"https://doi.org/{pub['doi']}"
                self._add_hyperlink(pub_para, doi_url, f"DOI: {pub['doi']}")
            elif pub.get("url"):
                pub_para.add_run(" ")
                self._add_hyperlink(pub_para, pub["url"], "[Link]")

    def _create_awards_section(self):
        """Creates the awards section."""
        awards = self.cv_data.data.get("awards", [])
        if not awards:
            return
        
        self._add_section_header("AWARDS & HONORS")
        for award in awards:
            award_para = self.doc.add_paragraph()
            self._set_paragraph_spacing(award_para, after=3)
            
            # Award name (bold)
            if award.get("name"):
                award_run = award_para.add_run(award["name"])
                self._set_font_properties(award_run, bold=True)
            
            # Issuer
            if award.get("issuer"):
                award_para.add_run(f", {award['issuer']}")
                for run in award_para.runs[-1:]:
                    self._set_font_properties(run)
            
            # Date
            if award.get("date"):
                date_str = self._format_date(award["date"])
                award_para.add_run(f", {date_str}")
                for run in award_para.runs[-1:]:
                    self._set_font_properties(run)
            
            # Description
            if award.get("description"):
                desc_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(desc_para, after=3)
                desc_run = desc_para.add_run(award["description"])
                self._set_font_properties(desc_run)

    def _create_volunteer_section(self):
        """Creates the volunteer work section."""
        volunteer_work = self.cv_data.data.get("volunteerWork", [])
        if not volunteer_work:
            return
        
        self._add_section_header("VOLUNTEER WORK")
        for work in volunteer_work:
            # Organization name
            if work.get("organization"):
                org_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(org_para)
                org_run = org_para.add_run(work["organization"])
                self._set_font_properties(org_run, bold=True)
            
            # Role
            if work.get("role"):
                role_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(role_para)
                role_run = role_para.add_run(work["role"])
                self._set_font_properties(role_run)
            
            # Dates
            start_date = self._format_date(work.get("startDate"))
            end_date = self._format_date(work.get("endDate"))
            if start_date or end_date:
                date_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(date_para)
                date_range = f"{start_date} - {end_date}" if start_date and end_date else start_date or end_date or ""
                date_run = date_para.add_run(date_range)
                self._set_font_properties(date_run)
            
            # Description
            if work.get("description"):
                self._add_bullet_point(work["description"])

    def _create_references_section(self):
        """Creates the references section."""
        references = self.cv_data.data.get("references", [])
        
        # By default, just add "References available upon request"
        # unless actual reference data is provided
        if not references:
            return
        
        self._add_section_header("REFERENCES")
        
        if len(references) == 0:
            ref_para = self.doc.add_paragraph()
            self._set_paragraph_spacing(ref_para)
            ref_run = ref_para.add_run("References available upon request.")
            self._set_font_properties(ref_run)
        else:
            # If references are explicitly provided, show them
            for ref in references:
                if ref.get("name"):
                    ref_para = self.doc.add_paragraph()
                    self._set_paragraph_spacing(ref_para)
                    ref_run = ref_para.add_run(ref["name"])
                    self._set_font_properties(ref_run, bold=True)
                    
                    # Add relationship and company on same line
                    ref_details = []
                    if ref.get("relationship"):
                        ref_details.append(ref["relationship"])
                    if ref.get("company"):
                        ref_details.append(ref["company"])
                    if ref_details:
                        ref_para.add_run(f", {', '.join(ref_details)}")
                        for run in ref_para.runs[-1:]:
                            self._set_font_properties(run)
                
                # Contact info on next line
                contact_parts = []
                if ref.get("email"):
                    contact_parts.append(ref["email"])
                if ref.get("phone"):
                    contact_parts.append(ref["phone"])
                if contact_parts:
                    contact_para = self.doc.add_paragraph()
                    self._set_paragraph_spacing(contact_para, after=3)
                    contact_run = contact_para.add_run(" | ".join(contact_parts))
                    self._set_font_properties(contact_run)

    def _add_section_header(self, text):
        header = self.doc.add_paragraph()
        self._set_paragraph_spacing(header, before=3, after=2)
        run = header.add_run(text)
        self._set_font_properties(run, bold=True)

    def _add_bullet_point(self, text):
        bullet_para = self.doc.add_paragraph(style='CustomBulletStyle')
        self._set_paragraph_spacing(bullet_para, after=3)
        run = bullet_para.add_run(text)
        self._set_font_properties(run)

    def _set_font_properties(self, run, bold=False):
        run.font.name = self.style["font_name"]
        run.font.size = Pt(self.style["font_size"])
        run.font.bold = bold

    def _set_paragraph_spacing(self, paragraph, before=None, after=None, line_spacing=None):
        p_spacing = self.style["paragraph_spacing"]
        paragraph.paragraph_format.space_before = Pt(before if before is not None else p_spacing["before"])
        paragraph.paragraph_format.space_after = Pt(after if after is not None else p_spacing["after"])
        paragraph.paragraph_format.line_spacing = line_spacing if line_spacing is not None else p_spacing["line_spacing"]

    def _add_hyperlink(self, paragraph, url, text, underline=True):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        
        if underline:
            rPr = OxmlElement('w:rPr')
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    @staticmethod
    def _format_date(date_str):
        """Formats date strings to human-readable format.
        
        Handles multiple date formats:
        - YYYY-MM-DD → January 15, 2025
        - YYYY-MM → January 2025
        - YYYY → 2025
        """
        if not date_str:
            return ""
        
        # Try different date formats
        date_formats = [
            ("%Y-%m-%d", "%B %d, %Y"),  # Full date
            ("%Y-%m", "%B %Y"),          # Year and month
            ("%Y", "%Y")                 # Year only
        ]
        
        for input_fmt, output_fmt in date_formats:
            try:
                return datetime.strptime(date_str, input_fmt).strftime(output_fmt)
            except ValueError:
                continue
        
        # If no format matches, return the original string
        return date_str

    @staticmethod
    def _extract_domain(url):
        if not url:
            return ""
        return url.replace("https://", "").replace("http://", "").split("/")[0]

    @staticmethod
    def _format_url(url):
        if not url:
            return ""
        
        domain = DocxGenerator._extract_domain(url)
        if "linkedin.com/in/" in url:
            try:
                username = url.split("/in/")[1].split("/")[0]
                return f"linkedin.com/in/{username}"
            except IndexError:
                return domain
        if "github.com/" in url:
            try:
                username = url.split("github.com/")[1].split("/")[0]
                return f"github.com/{username}"
            except IndexError:
                return domain
        return domain

    @staticmethod
    def _find_company_url(job_data):
        for field in ["companyUrl", "company_url", "website", "url", "link", "homepage"]:
            if job_data.get(field):
                return job_data[field]
        return None

    @staticmethod
    def _format_language_entry(language):
        name = language.get("language")
        if not name:
            return ""
        if language.get("native") or language.get("proficiency") == "C2":
            return f"{name} (native)"
        if language.get("proficiency"):
            return f"{name} ({language.get('proficiency')})"
        return name

def main():
    """Backwards compatibility wrapper - redirects to new CLI."""
    # Import here to avoid circular imports
    from .__main__ import main as cvac_main
    
    # Convert old-style arguments to new CLI format
    if len(sys.argv) < 2:
        print("Usage: python cv-to-docx.py <path_to_cv.json> [output_path.docx]")
        print("\nNote: This script is deprecated. Please use the new CLI:")
        print("  cvac generate <input> <output>")
        sys.exit(1)
    
    # Prepare arguments for new CLI
    new_args = ['cvac', 'generate']
    new_args.extend(sys.argv[1:])
    
    # Replace sys.argv temporarily
    old_argv = sys.argv
    sys.argv = new_args
    
    try:
        return cvac_main()
    finally:
        sys.argv = old_argv

if __name__ == "__main__":
    main()
