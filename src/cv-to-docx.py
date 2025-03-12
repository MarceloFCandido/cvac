import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Twips
from docx.enum.style import WD_STYLE_TYPE

def format_date(date_str):
    """Convert YYYY-MM format to Month YYYY"""
    if not date_str:
        return ""
    
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m")
        return date_obj.strftime("%B %Y")
    except ValueError:
        return date_str

def add_hyperlink(paragraph, url, text, color=None, underline=True):
    """
    A function that places a hyperlink within a paragraph object.
    
    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run element
    new_run = OxmlElement('w:r')
    
    # If underline is enabled, add underlining
    if underline:
        rPr = OxmlElement('w:rPr')
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        new_run.append(rPr)

    # Create a text element
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    # Add the run to the hyperlink
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink

def set_font_properties(run, font_name="Calibri", size=11, bold=False):
    """Set consistent font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold

def add_quote_character(paragraph):
    """Add a quote character to the beginning of a paragraph (empty function now)"""
    # No longer adding the '>' character
    return

def set_paragraph_spacing(paragraph, before=0, after=2, line_spacing=1):
    """Set consistent paragraph spacing"""
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing

def create_custom_bullet_style(doc):
    """Create the custom bullet style that matches the target document"""
    # Create a new list style
    styles = doc.styles
    list_style = styles.add_style('CustomBulletStyle', WD_STYLE_TYPE.PARAGRAPH)
    list_style.base_style = styles['List Bullet']
    
    # Customize the style to match the target document
    list_style.paragraph_format.left_indent = Cm(0.5)
    list_style.paragraph_format.first_line_indent = Cm(-0.25)
    list_style.paragraph_format.line_spacing = 1.0
    list_style.paragraph_format.space_after = Pt(3)
    list_style.paragraph_format.space_before = Pt(0)
    
    # Set the font for the style
    list_style.font.name = "Calibri"
    list_style.font.size = Pt(11)
    
    return list_style

def extract_domain_from_url(url):
    """Extract domain from URL for displaying in the document"""
    if not url:
        return ""
    
    # Remove protocol prefix
    domain = url.replace("https://", "").replace("http://", "")
    
    # Remove trailing slash and anything after it
    if "/" in domain:
        domain = domain.split("/")[0]
        
    return domain

def find_company_url(job_data):
    """Find a company URL in the job data, checking multiple possible fields"""
    url_field_candidates = [
        "companyUrl", "company_url", "website", "url", "link", "homepage"
    ]
    
    for field in url_field_candidates:
        if field in job_data and job_data[field]:
            return job_data[field]
    
    return None

def format_language_entry(language):
    """Format a language entry based on available data"""
    language_name = language.get("language", "")
    proficiency = language.get("proficiency", "")
    native = language.get("native", False)
    
    if not language_name:
        return ""
    
    # Check if this is marked as native language
    if native:
        return f"{language_name} (native)"
    
    # Check if proficiency level C2 is present (often considered native-like)
    if proficiency == "C2":
        return f"{language_name} (native)"
    
    # Otherwise format with proficiency if available
    if proficiency:
        return f"{language_name} ({proficiency})"
    
    return language_name

def create_cv_document(json_path, output_path="CV.docx"):
    """Create a Word document from JSON CV data that matches the provided template"""
    
    # Load JSON data
    with open(json_path, 'r', encoding='utf-8') as file:
        cv_data = json.load(file)
    
    # Create document with exact formatting to match the reference
    doc = Document()
    
    # Set exact margins to match the target document - using smaller margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
    
    # Set default font for the document
    styles = doc.styles
    style = styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Create custom bullet style to match target document
    bullet_style = create_custom_bullet_style(doc)
    
    # Personal Information section
    personal_info = cv_data.get("personalInfo", {})
    if personal_info:
        # Full Name
        name_parts = []
        for name_field in ["firstName", "middleName", "lastName"]:
            if name_field in personal_info and personal_info[name_field]:
                name_parts.append(personal_info[name_field])
        
        full_name = " ".join(name_parts)
        
        if full_name:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(full_name)
            name_run.font.size = Pt(14)
            name_run.font.name = "Calibri"
            name_run.font.bold = True
            set_paragraph_spacing(name_para, before=0, after=3)
        
        # Contact Info - First row
        contact_row1_items = []
        if "email" in personal_info:
            contact_row1_items.append({
                "text": personal_info["email"],
                "url": f"mailto:{personal_info['email']}"
            })
        
        if "phone" in personal_info:
            contact_row1_items.append({
                "text": personal_info["phone"],
                "url": f"tel:{personal_info['phone']}"
            })
        
        if contact_row1_items:
            contact_para1 = doc.add_paragraph()
            contact_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(contact_para1, before=0, after=2)
            
            for i, item in enumerate(contact_row1_items):
                if i > 0:
                    contact_para1.add_run(" | ")
                
                add_hyperlink(contact_para1, item["url"], item["text"], underline=True)
            
            for run in contact_para1.runs:
                set_font_properties(run)
        
        # Contact Info - Second row
        contact_row2_items = []
        
        # LinkedIn
        if "linkedIn" in personal_info and personal_info["linkedIn"]:
            linkedin_url = personal_info["linkedIn"]
            # Include www. if needed based on convention
            if "www." not in linkedin_url and linkedin_url.startswith("https://"):
                linkedin_url = linkedin_url.replace("https://", "https://www.")
            
            linkedin_text = extract_domain_from_url(linkedin_url)
            if "linkedin.com" in linkedin_url.lower():
                # Extract username from LinkedIn URL if possible
                if "/in/" in linkedin_url.lower():
                    username = linkedin_url.lower().split("/in/")[1].split("/")[0].split("?")[0]
                    linkedin_text = f"linkedin.com/in/{username}"
            
            contact_row2_items.append({
                "text": linkedin_text,
                "url": linkedin_url
            })
        
        # GitHub
        if "githubUrl" in personal_info and personal_info["githubUrl"]:
            github_url = personal_info["githubUrl"]
            github_text = extract_domain_from_url(github_url)
            
            if "github.com" in github_url.lower():
                # Extract username from GitHub URL if possible
                username = github_url.split("github.com/")[1].split("/")[0]
                github_text = f"github.com/{username}"
            
            contact_row2_items.append({
                "text": github_text,
                "url": github_url
            })
        
        # Website or other contact methods
        for field in ["website", "portfolio", "blog"]:
            if field in personal_info and personal_info[field]:
                url = personal_info[field]
                text = extract_domain_from_url(url)
                contact_row2_items.append({
                    "text": text,
                    "url": url
                })
        
        # Add contact row 2 if items exist
        if contact_row2_items:
            contact_para2 = doc.add_paragraph()
            contact_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(contact_para2, before=0, after=3)
            
            for i, item in enumerate(contact_row2_items):
                if i > 0:
                    contact_para2.add_run(" | ")
                
                add_hyperlink(contact_para2, item["url"], item["text"], underline=True)
            
            for run in contact_para2.runs:
                set_font_properties(run)
    
    # EXPERIENCE Section
    work_experience = cv_data.get("workExperience", [])
    if work_experience:
        # Experience header
        exp_header = doc.add_paragraph()
        exp_run = exp_header.add_run("EXPERIENCE")
        set_font_properties(exp_run, bold=True)
        set_paragraph_spacing(exp_header, before=3, after=2)
        
        # Work Experience entries
        for job in work_experience:
            # Company name
            if "company" in job:
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(job["company"])
                set_font_properties(company_run, bold=True)
                set_paragraph_spacing(company_para, before=0, after=2)
                
                # Find and add company URL if available
                company_url = find_company_url(job)
                if company_url:
                    company_domain = extract_domain_from_url(company_url)
                    company_para.add_run(" (")
                    add_hyperlink(company_para, company_url, company_domain, underline=False)
                    company_para.add_run(")")
                    for run in company_para.runs:
                        set_font_properties(run)
            
            # Position
            if "position" in job:
                position_para = doc.add_paragraph()
                position_run = position_para.add_run(job["position"])
                set_font_properties(position_run, bold=True)
                set_paragraph_spacing(position_para, before=0, after=2)
            
            # Date range with single dash
            start_date = format_date(job.get("startDate", ""))
            end_date = "Present" if job.get("current", False) else format_date(job.get("endDate", ""))
            
            if start_date or end_date:
                date_para = doc.add_paragraph()
                date_range = ""
                
                if start_date and end_date:
                    date_range = f"{start_date} - {end_date}"
                elif start_date:
                    date_range = start_date
                elif end_date:
                    date_range = end_date
                
                date_run = date_para.add_run(date_range)
                set_font_properties(date_run)
                set_paragraph_spacing(date_para, before=0, after=2)
            
            # Company description as bullet point
            if "description" in job and job["description"]:
                bullet_para = doc.add_paragraph(style='CustomBulletStyle')
                bullet_para.style = bullet_style
                set_paragraph_spacing(bullet_para, before=0, after=3)
                description_run = bullet_para.add_run(job["description"])
                set_font_properties(description_run)
            
            # Achievements as bullet points
            achievements = job.get("achievements", [])
            
            # Process each achievement as a bullet point
            for achievement in achievements:
                if not achievement:
                    continue
                
                # Ensure periods at the end of each bullet point for consistency
                if not achievement.endswith('.'):
                    achievement += '.'
                
                bullet_para = doc.add_paragraph(style='CustomBulletStyle')
                bullet_para.style = bullet_style
                achievement_run = bullet_para.add_run(achievement)
                set_font_properties(achievement_run)
                set_paragraph_spacing(bullet_para, before=0, after=3)
    
    # EDUCATION Section
    education = cv_data.get("education", [])
    if education:
        # Education header
        edu_header = doc.add_paragraph()
        edu_run = edu_header.add_run("EDUCATION")
        set_font_properties(edu_run, bold=True)
        set_paragraph_spacing(edu_header, before=3, after=2)
        
        # Education entries
        for edu in education:
            # Institution
            if "institution" in edu and edu["institution"]:
                inst_para = doc.add_paragraph()
                inst_run = inst_para.add_run(edu["institution"])
                set_font_properties(inst_run)
                set_paragraph_spacing(inst_para, before=0, after=2)
            
            # Degree and field
            degree_parts = []
            if "degree" in edu and edu["degree"]:
                degree_parts.append(edu["degree"])
            if "field" in edu and edu["field"]:
                degree_parts.append(edu["field"])
            
            if degree_parts:
                degree_para = doc.add_paragraph()
                degree_run = degree_para.add_run(" | ".join(degree_parts))
                set_font_properties(degree_run)
                set_paragraph_spacing(degree_para, before=0, after=2)
            
            # Location and graduation date
            location_grad_parts = []
            if "location" in edu and edu["location"]:
                location_grad_parts.append(edu["location"])
            if "graduationDate" in edu and edu["graduationDate"]:
                location_grad_parts.append(f"Graduated in {format_date(edu['graduationDate'])}")
            
            if location_grad_parts:
                loc_date_para = doc.add_paragraph()
                loc_date_run = loc_date_para.add_run(" | ".join(location_grad_parts))
                set_font_properties(loc_date_run)
                set_paragraph_spacing(loc_date_para, before=0, after=3)
    
    # LANGUAGES Section
    languages = cv_data.get("languages", [])
    if languages:
        # Languages header
        lang_header = doc.add_paragraph()
        lang_run = lang_header.add_run("LANGUAGES")
        set_font_properties(lang_run, bold=True)
        set_paragraph_spacing(lang_header, before=3, after=2)
        
        # Format and add each language
        lang_entries = []
        
        for lang in languages:
            formatted_lang = format_language_entry(lang)
            if formatted_lang:
                lang_entries.append(formatted_lang)
        
        # Join all language entries with commas
        if lang_entries:
            lang_para = doc.add_paragraph()
            lang_run = lang_para.add_run(", ".join(lang_entries))
            set_font_properties(lang_run)
            set_paragraph_spacing(lang_para, before=0, after=3)
    
    # TECHNOLOGIES/SKILLS Section
    skills = cv_data.get("skills", [])
    if skills:
        # Technologies header
        tech_header = doc.add_paragraph()
        tech_run = tech_header.add_run("TECHNOLOGIES")
        set_font_properties(tech_run, bold=True)
        set_paragraph_spacing(tech_header, before=3, after=2)
        
        # Extract skill names, handling different potential formats
        skill_names = []
        for skill in skills:
            # Handle dict format with "name" key
            if isinstance(skill, dict) and "name" in skill:
                skill_names.append(skill["name"])
            # Handle string format
            elif isinstance(skill, str):
                skill_names.append(skill)
            # Handle dict format with potential alternative keys
            elif isinstance(skill, dict):
                for key in ["skill", "technology", "tech", "title"]:
                    if key in skill:
                        skill_names.append(skill[key])
                        break
        
        # Format skills list with period at end
        if skill_names:
            skills_text = ", ".join(skill_names)
            if not skills_text.endswith("."):
                skills_text += "."
            
            tech_para = doc.add_paragraph()
            tech_run = tech_para.add_run(skills_text)
            set_font_properties(tech_run)
    
    # Save the document
    doc.save(output_path)
    print(f"CV saved as {output_path}")
    return output_path

if __name__ == "__main__":
    # Get input and output filenames
    import sys
    
    if len(sys.argv) >= 2:
        json_file = sys.argv[1]
    else:
        json_file = input("Enter the path to your JSON CV file: ")
    
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        output_file = "resume-generated.docx"
    
    if os.path.exists(json_file):
        create_cv_document(json_file, output_file)
    else:
        print(f"Error: {json_file} not found. Please ensure the file exists in the specified path.")
