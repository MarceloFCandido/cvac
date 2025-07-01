import json
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import jsonschema
from style import STYLE_CONFIG

class CVData:
    """Handles loading, validation, and access to CV data."""
    def __init__(self, json_path, schema_path):
        self.json_path = json_path
        self.schema_path = schema_path
        self.data = self._load_and_validate_data()

    def _load_and_validate_data(self):
        """Loads the JSON data and validates it against the schema."""
        try:
            with open(self.json_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            with open(self.schema_path, 'r', encoding='utf-8') as file:
                schema = json.load(file)
            
            jsonschema.validate(instance=data, schema=schema)
            return data
        except FileNotFoundError as e:
            print(f"Error: {e.filename} not found.")
            sys.exit(1)
        except jsonschema.exceptions.ValidationError as e:
            print(f"JSON validation error: {e.message}")
            sys.exit(1)
        except json.JSONDecodeError:
            print(f"Error decoding JSON from {self.json_path}")
            sys.exit(1)

class DocxGenerator:
    """Generates the CV document from CVData."""
    def __init__(self, cv_data, style_config):
        self.cv_data = cv_data
        self.style = style_config
        self.doc = Document()
        self._apply_document_styles()

    def _apply_document_styles(self):
        """Applies base styles and margins to the document."""
        for section in self.doc.sections:
            section.top_margin = self.style["margins"]["top"]
            section.bottom_margin = self.style["margins"]["bottom"]
            section.left_margin = self.style["margins"]["left"]
            section.right_margin = self.style["margins"]["right"]

        style = self.doc.styles['Normal']
        style.font.name = self.style["font_name"]
        style.font.size = Pt(self.style["font_size"])
        
        self._create_custom_bullet_style()

    def _create_custom_bullet_style(self):
        """Creates the custom bullet style."""
        bullet_style_def = self.style["bullet_style"]
        list_style = self.doc.styles.add_style('CustomBulletStyle', WD_STYLE_TYPE.PARAGRAPH)
        list_style.base_style = self.doc.styles['List Bullet']
        list_style.paragraph_format.left_indent = bullet_style_def["left_indent"]
        list_style.paragraph_format.first_line_indent = bullet_style_def["first_line_indent"]
        list_style.paragraph_format.line_spacing = bullet_style_def["line_spacing"]
        list_style.paragraph_format.space_after = bullet_style_def["space_after"]
        list_style.paragraph_format.space_before = bullet_style_def["space_before"]
        list_style.font.name = self.style["font_name"]
        list_style.font.size = Pt(self.style["font_size"])

    def generate(self, output_path):
        """Generates and saves the full CV document."""
        self._create_personal_info_section()
        self._create_experience_section()
        self._create_education_section()
        self._create_languages_section()
        self._create_skills_section()
        self.doc.save(output_path)
        print(f"CV saved as {output_path}")

    def _create_personal_info_section(self):
        personal_info = self.cv_data.data.get("personalInfo", {})
        if not personal_info:
            return

        # Full Name
        name_parts = [personal_info.get(key) for key in ["firstName", "middleName", "lastName"] if personal_info.get(key)]
        full_name = " ".join(name_parts)
        if full_name:
            name_para = self.doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(full_name)
            name_run.font.size = Pt(self.style["name_font_size"])
            name_run.font.name = self.style["font_name"]
            name_run.font.bold = True
            self._set_paragraph_spacing(name_para, after=3)

        # Contact Info Rows
        self._add_contact_row([
            {"text": personal_info.get("email"), "url": f"mailto:{personal_info.get('email')}"},
            {"text": personal_info.get("phone"), "url": f"tel:{personal_info.get('phone')}"}
        ])
        
        contact_row2_items = []
        if personal_info.get("linkedIn"):
            contact_row2_items.append({"text": self._format_url(personal_info["linkedIn"]), "url": personal_info["linkedIn"]})
        if personal_info.get("githubUrl"):
            contact_row2_items.append({"text": self._format_url(personal_info["githubUrl"]), "url": personal_info["githubUrl"]})
        for field in ["website", "portfolio", "blog"]:
            if personal_info.get(field):
                contact_row2_items.append({"text": self._extract_domain(personal_info[field]), "url": personal_info[field]})
        self._add_contact_row(contact_row2_items, after=3)

    def _add_contact_row(self, items, after=2):
        items = [item for item in items if item.get("text")]
        if not items:
            return
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_spacing(para, after=after)
        
        for i, item in enumerate(items):
            if i > 0:
                para.add_run(" | ")
            self._add_hyperlink(para, item["url"], item["text"])
        
        for run in para.runs:
            self._set_font_properties(run)

    def _create_experience_section(self):
        work_experience = self.cv_data.data.get("workExperience", [])
        if not work_experience:
            return
        
        self._add_section_header("EXPERIENCE")
        for job in work_experience:
            if job.get("company"):
                company_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(company_para)
                company_run = company_para.add_run(job["company"])
                self._set_font_properties(company_run, bold=True)
                
                company_url = self._find_company_url(job)
                if company_url:
                    company_para.add_run(" (")
                    self._add_hyperlink(company_para, company_url, self._extract_domain(company_url), underline=False)
                    company_para.add_run(")")
                    for run in company_para.runs:
                        self._set_font_properties(run, bold=True)

            if job.get("position"):
                position_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(position_para)
                position_run = position_para.add_run(job["position"])
                self._set_font_properties(position_run, bold=True)

            start_date = self._format_date(job.get("startDate"))
            end_date = "Present" if job.get("current") else self._format_date(job.get("endDate"))
            if start_date or end_date:
                date_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(date_para)
                date_range = f"{start_date} - {end_date}" if start_date and end_date else start_date or end_date
                date_run = date_para.add_run(date_range)
                self._set_font_properties(date_run)

            if job.get("description"):
                self._add_bullet_point(job["description"])

            for achievement in job.get("achievements", []):
                if achievement:
                    self._add_bullet_point(achievement if achievement.endswith('.') else f"{achievement}.")

    def _create_education_section(self):
        education = self.cv_data.data.get("education", [])
        if not education:
            return

        self._add_section_header("EDUCATION")
        for edu in education:
            if edu.get("institution"):
                inst_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(inst_para)
                inst_run = inst_para.add_run(edu["institution"])
                self._set_font_properties(inst_run)

            degree_parts = [edu.get("degree"), edu.get("field")]
            if any(degree_parts):
                degree_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(degree_para)
                degree_run = degree_para.add_run(" | ".join(filter(None, degree_parts)))
                self._set_font_properties(degree_run)

            loc_grad_parts = [edu.get("location")]
            if edu.get("graduationDate"):
                loc_grad_parts.append(f"Graduated in {self._format_date(edu['graduationDate'])}")
            if any(loc_grad_parts):
                loc_date_para = self.doc.add_paragraph()
                self._set_paragraph_spacing(loc_date_para, after=3)
                loc_date_run = loc_date_para.add_run(" | ".join(filter(None, loc_grad_parts)))
                self._set_font_properties(loc_date_run)

    def _create_languages_section(self):
        languages = self.cv_data.data.get("languages", [])
        if not languages:
            return
        
        self._add_section_header("LANGUAGES")
        lang_entries = [self._format_language_entry(lang) for lang in languages]
        if lang_entries:
            lang_para = self.doc.add_paragraph()
            self._set_paragraph_spacing(lang_para, after=3)
            lang_run = lang_para.add_run(", ".join(filter(None, lang_entries)))
            self._set_font_properties(lang_run)

    def _create_skills_section(self):
        skills = self.cv_data.data.get("skills", [])
        if not skills:
            return
        
        self._add_section_header("TECHNOLOGIES")
        skill_names = []
        for skill in skills:
            if isinstance(skill, dict):
                skill_names.append(skill.get("name"))
            elif isinstance(skill, str):
                skill_names.append(skill)
        
        skill_names = list(filter(None, skill_names))
        if skill_names:
            skills_text = ", ".join(skill_names)
            if not skills_text.endswith("."):
                skills_text += "."
            
            tech_para = self.doc.add_paragraph()
            tech_run = tech_para.add_run(skills_text)
            self._set_font_properties(tech_run)

    def _add_section_header(self, text):
        header = self.doc.add_paragraph()
        self._set_paragraph_spacing(header, before=3, after=2)
        run = header.add_run(text)
        self._set_font_properties(run, bold=True)

    def _add_bullet_point(self, text):
        bullet_para = self.doc.add_paragraph(style='CustomBulletStyle')
        self._set_paragraph_spacing(bullet_para, after=3)
        run = bullet_para.add_run(text)
        self._set_font_properties(run)

    def _set_font_properties(self, run, bold=False):
        run.font.name = self.style["font_name"]
        run.font.size = Pt(self.style["font_size"])
        run.font.bold = bold

    def _set_paragraph_spacing(self, paragraph, before=None, after=None, line_spacing=None):
        p_spacing = self.style["paragraph_spacing"]
        paragraph.paragraph_format.space_before = Pt(before if before is not None else p_spacing["before"])
        paragraph.paragraph_format.space_after = Pt(after if after is not None else p_spacing["after"])
        paragraph.paragraph_format.line_spacing = line_spacing if line_spacing is not None else p_spacing["line_spacing"]

    def _add_hyperlink(self, paragraph, url, text, underline=True):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        
        if underline:
            rPr = OxmlElement('w:rPr')
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    @staticmethod
    def _format_date(date_str):
        if not date_str:
            return ""
        try:
            return datetime.strptime(date_str, "%Y-%m").strftime("%B %Y")
        except ValueError:
            return date_str

    @staticmethod
    def _extract_domain(url):
        if not url:
            return ""
        return url.replace("https://", "").replace("http://", "").split("/")[0]

    @staticmethod
    def _format_url(url):
        if not url:
            return ""
        
        domain = DocxGenerator._extract_domain(url)
        if "linkedin.com/in/" in url:
            try:
                username = url.split("/in/")[1].split("/")[0]
                return f"linkedin.com/in/{username}"
            except IndexError:
                return domain
        if "github.com/" in url:
            try:
                username = url.split("github.com/")[1].split("/")[0]
                return f"github.com/{username}"
            except IndexError:
                return domain
        return domain

    @staticmethod
    def _find_company_url(job_data):
        for field in ["companyUrl", "company_url", "website", "url", "link", "homepage"]:
            if job_data.get(field):
                return job_data[field]
        return None

    @staticmethod
    def _format_language_entry(language):
        name = language.get("language")
        if not name:
            return ""
        if language.get("native") or language.get("proficiency") == "C2":
            return f"{name} (native)"
        if language.get("proficiency"):
            return f"{name} ({language.get('proficiency')})"
        return name

def main():
    if len(sys.argv) < 2:
        print("Usage: python cv-to-docx.py <path_to_cv.json> [output_path.docx]")
        sys.exit(1)

    json_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else "resume-generated.docx"
    
    # Assume schema is in the `schema` directory, relative to the project root.
    script_dir = os.path.dirname(os.path.abspath(__file__))
    # Go up one level from `src` to the project root
    project_root = os.path.dirname(script_dir)
    schema_file = os.path.join(project_root, 'schema', 'cv.schema.json')


    cv_data = CVData(json_file, schema_file)
    generator = DocxGenerator(cv_data, STYLE_CONFIG)
    generator.generate(output_file)

if __name__ == "__main__":
    main()
