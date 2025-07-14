import pytest
import os
import json
import sys

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from docx import Document
from cvac.cv_to_docx import CVData, DocxGenerator
from cvac.core.style_loader import StyleLoader

@pytest.fixture
def cv_data_instance(tmp_path):
    # Create a dummy schema and cv.json for testing
    schema_content = {
        "$schema": "http://json-schema.org/draft-07/schema#",
        "title": "CV Schema",
        "type": "object",
        "required": ["personalInfo", "workExperience", "education"],
        "properties": {
            "personalInfo": {
                "type": "object",
                "required": ["firstName", "lastName", "email"],
                "properties": {
                    "firstName": {"type": "string"},
                    "lastName": {"type": "string"},
                    "email": {"type": "string", "format": "email"}
                }
            },
            "workExperience": {"type": "array"},
            "education": {"type": "array"}
        }
    }
    cv_content = {
        "personalInfo": {
            "firstName": "Test",
            "lastName": "User",
            "email": "test.user@example.com"
        },
        "workExperience": [],
        "education": []
    }
    schema_file = tmp_path / "schema.json"
    schema_file.write_text(json.dumps(schema_content))
    cv_file = tmp_path / "cv.json"
    cv_file.write_text(json.dumps(cv_content))
    return CVData(str(cv_file), str(schema_file))

def test_document_generation(cv_data_instance, tmp_path):
    output_path = tmp_path / "test_cv.docx"
    # Use StyleLoader to get properly converted style config
    style_loader = StyleLoader()
    style_config = style_loader.load_style()  # This will convert units properly
    generator = DocxGenerator(cv_data_instance, style_config)
    generator.generate(str(output_path))

    assert os.path.exists(output_path)

    doc = Document(output_path)
    assert len(doc.paragraphs) > 0

    # Check for the name in the document
    name_found = any("Test User" in p.text for p in doc.paragraphs)
    assert name_found, "Name 'Test User' not found in the generated document."
